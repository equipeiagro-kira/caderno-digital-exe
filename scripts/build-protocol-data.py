"""Build the static research-protocol dataset used by the offline PWA."""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parents[1]
WORKBOOK = Path(r"C:\Users\Lenovo\Desktop\TRATAMENTOS - PROTOCOLOS.xlsx")
FIELD_BOOK = Path(r"C:\Users\Lenovo\Desktop\CADERNO DE CAMPO.docx")
OUTPUT = ROOT / "src" / "protocols-data.js"


PROTOCOLS = [
    {"id": "bayer", "nome": "Bayer Fungicidas", "short": "BAYER", "title_row": 1, "header_row": 2, "end_row": 9, "tables": [1, 2, 3, 4, 5], "measurements": ["Fitotoxicidade", "Severidade de doencas"]},
    {"id": "corteva-fungicidas", "nome": "Corteva Fungicidas", "short": "CTVA FUNGI.", "title_row": 11, "header_row": 12, "end_row": 20, "tables": [6, 7, 8, 9, 10], "measurements": ["Severidade e desfolha", "Fitotoxicidade"]},
    {"id": "excelencia-fungicidas", "nome": "Excelencia Fungicidas", "short": "EXC. FUNGI", "title_row": 22, "header_row": 23, "end_row": 32, "tables": [11, 12, 13, 14, 15], "measurements": ["Severidade de doencas", "Fitotoxicidade"]},
    {"id": "koppert", "nome": "Koppert", "short": "KOPPERT", "title_row": 34, "header_row": 35, "end_row": 41, "tables": [16, 17, 18, 19, 20], "measurements": ["Severidade de doencas", "Fitotoxicidade"]},
    {"id": "fmc", "nome": "FMC", "short": "FMC", "title_row": 43, "header_row": 44, "end_row": 50, "tables": [21, 22, 23, 24, 25], "measurements": ["Severidade e desfolha", "Fitotoxicidade"]},
    {"id": "simbiose", "nome": "Simbiose Fungicida", "short": "SIMBIOSE FUNGI.", "title_row": 52, "header_row": 53, "end_row": 57, "tables": [26, 27, 28, 29, 30], "measurements": ["Severidade e desfolha", "Fitotoxicidade"]},
    {"id": "c6-bio-cobre", "nome": "C6 Bio - Cobre", "short": "C6 BIO COBRE", "title_row": 59, "header_row": 60, "end_row": 65, "tables": [31, 32, 33, 34, 35], "measurements": ["Severidade e controle", "Fitotoxicidade"]},
    {"id": "verdivo", "nome": "Verdivo", "short": "VERDIVO", "title_row": 67, "header_row": 68, "end_row": 75, "tables": [36, 37, 38, 39, 40], "measurements": ["Severidade e desfolha", "Fitotoxicidade"]},
    {"id": "syngenta", "nome": "Syngenta", "short": "SYNGENTA", "title_row": 77, "header_row": 78, "end_row": 84, "tables": [41, 42, 43, 44, 45], "measurements": ["Severidade e desfolha", "Fitotoxicidade"]},
    {"id": "upl", "nome": "UPL", "short": "UPL", "title_row": 86, "header_row": 87, "end_row": 93, "tables": [46, 47, 48, 49, 50], "measurements": ["Severidade e desfolha", "Fitotoxicidade"]},
    {"id": "icl", "nome": "ICL", "short": "ICL", "title_row": 95, "header_row": 96, "end_row": 101, "tables": [51, 52, 53, 54, 55], "measurements": ["Severidade e controle", "Fitotoxicidade"]},
    {"id": "corteva-utrisha", "nome": "Corteva - Utrisha N", "short": "CTVA UTRISHA", "title_row": 103, "header_row": 104, "end_row": 106, "tables": [56, 57, 58], "measurements": []},
    {"id": "nitro", "nome": "Nitro", "short": "NITRO", "title_row": 108, "header_row": 109, "end_row": 114, "tables": [59, 60, 61, 62], "measurements": ["Fitotoxicidade"]},
    {"id": "multitecnica", "nome": "Multitecnica", "short": "MULTITEC", "title_row": 121, "header_row": 122, "end_row": 127, "tables": [63, 64, 65], "measurements": []},
    {"id": "bio-azon", "nome": "Bio Azon", "short": "BIO AZON", "title_row": 129, "header_row": 130, "end_row": 135, "tables": [], "measurements": []},
]


GENERAL_MAP = [
    {"id": "corteva-fungicidas", "label": "CTVA FUNGI.", "col": 1, "row": 1, "span": 8, "tone": "cyan"},
    {"id": "bayer", "label": "BAYER FUNGI.", "col": 1, "row": 9, "span": 7, "tone": "orange"},
    {"id": "koppert", "label": "KOPPERT", "col": 2, "row": 1, "span": 6, "tone": "orange"},
    {"id": "excelencia-fungicidas", "label": "EXC. FUNGI", "col": 2, "row": 7, "span": 9, "tone": "cyan"},
    {"id": "c6-bio-cobre", "label": "C6 BIO COBRE", "col": 3, "row": 1, "span": 5, "tone": "yellow"},
    {"id": "simbiose", "label": "SIMBIOSE FUNGI.", "col": 3, "row": 6, "span": 4, "tone": "cyan"},
    {"id": "fmc", "label": "FMC", "col": 3, "row": 10, "span": 6, "tone": "cyan"},
    {"id": "unidentified", "label": "AREA NAO IDENTIFICADA", "sourceLabel": "XXXXX", "col": 4, "row": 1, "span": 2, "tone": "muted"},
    {"id": "syngenta", "label": "SYNGENTA", "col": 4, "row": 3, "span": 6, "tone": "cyan"},
    {"id": "verdivo", "label": "VERDIVO", "col": 4, "row": 9, "span": 7, "tone": "cyan"},
    {"id": "corteva-utrisha", "label": "CTVA UTRISHA", "col": 5, "row": 1, "span": 4, "tone": "yellow"},
    {"id": "icl", "label": "ICL", "col": 5, "row": 5, "span": 5, "tone": "orange"},
    {"id": "upl", "label": "UPL", "col": 5, "row": 10, "span": 6, "tone": "cyan"},
    {"id": "bio-azon", "label": "BIO AZON", "col": 6, "row": 1, "span": 5, "tone": "yellow"},
    {"id": "multitecnica", "label": "MULTITEC", "col": 6, "row": 6, "span": 5, "tone": "orange"},
    {"id": "nitro", "label": "NITRO", "col": 6, "row": 11, "span": 5, "tone": "cyan"},
]


def clean(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        value = int(value)
    return re.sub(r"\s+", " ", str(value)).strip()


def table_rows(table) -> list[list[str]]:
    return [[clean(cell.text) for cell in row.cells] for row in table.rows]


def extract_treatments(ws, spec: dict[str, Any]) -> tuple[list[str], list[dict[str, Any]]]:
    headers = [clean(ws.cell(spec["header_row"], col).value) for col in range(2, 7)]
    headers = [header for header in headers if header]
    treatments = []
    for row in range(spec["header_row"] + 1, spec["end_row"] + 1):
        number = clean(ws.cell(row, 1).value)
        if not number:
            continue
        applications = []
        for offset, stage in enumerate(headers, start=2):
            recipe = clean(ws.cell(row, offset).value)
            if recipe:
                applications.append({"epoca": stage, "receita": recipe})
        treatments.append({"numero": number, "aplicacoes": applications})
    return headers, treatments


def extract_randomization(table) -> list[dict[str, Any]]:
    rows = table_rows(table)
    logical_rows = rows[:-1:2]
    columns = list(zip(*[row[:4] for row in logical_rows])) if logical_rows else []
    return [
        {"repeticao": f"R{index + 1}", "parcelas": [clean(value) for value in column]}
        for index, column in enumerate(columns)
    ]


def extract_evaluations(table) -> list[dict[str, Any]]:
    rows = table_rows(table)
    if not rows:
        return []
    width = min(5, len(rows[0]))
    evaluations = []
    for col in range(width):
        items = []
        for row in rows[1:]:
            value = clean(row[col]) if col < len(row) else ""
            if value and value not in items:
                items.append(value)
        evaluations.append({"fase": clean(rows[0][col]), "itens": items})
    return evaluations


def extract_data_table(table, title: str) -> dict[str, Any]:
    rows = table_rows(table)
    if not rows:
        return {"titulo": title, "colunas": [], "linhas": []}
    width = len(rows[0])
    return {
        "titulo": title,
        "colunas": rows[0],
        "linhas": [(row + [""] * width)[:width] for row in rows[1:]],
    }


def extract_delta_composition(ws) -> list[dict[str, str]]:
    stages = [clean(ws.cell(116, col).value) for col in range(2, 6)]
    result = []
    for col, stage in enumerate(stages, start=2):
        parts = [clean(ws.cell(row, col).value).strip(" +") for row in range(117, 120)]
        recipe = " + ".join(part for part in parts if part)
        if stage and recipe:
            result.append({"epoca": stage, "receita": recipe})
    return result


def main() -> None:
    ws = load_workbook(WORKBOOK, data_only=True, read_only=True)["TRATAMENTOS"]
    document = Document(FIELD_BOOK)
    protocols = []

    for spec in PROTOCOLS:
        stages, treatments = extract_treatments(ws, spec)
        protocol = {
            "id": spec["id"],
            "nome": spec["nome"],
            "nomeFonte": clean(ws.cell(spec["title_row"], 1).value),
            "nomeCurto": spec["short"],
            "epocas": stages,
            "totalTratamentos": len(treatments),
            "tratamentos": treatments,
            "croqui": [],
            "avaliacoes": [],
            "campo": {"estande": None, "medicoes": []},
            "avisos": [],
        }

        if spec["tables"]:
            croqui_table, evaluation_table, stand_table, *measurement_tables = spec["tables"]
            protocol["croqui"] = extract_randomization(document.tables[croqui_table])
            protocol["avaliacoes"] = extract_evaluations(document.tables[evaluation_table])
            protocol["campo"]["estande"] = extract_data_table(document.tables[stand_table], "Estande")
            protocol["campo"]["medicoes"] = [
                extract_data_table(document.tables[index], title)
                for index, title in zip(measurement_tables, spec["measurements"])
            ]
        else:
            protocol["avisos"].append("A ficha e o croqui deste protocolo nao constam no caderno de campo recebido.")

        if spec["id"] == "corteva-utrisha":
            protocol["avisos"].append(
                "Divergencia de origem: o croqui mostra tratamentos 1 a 4, mas a planilha de tratamentos e o estande possuem apenas 1 e 2."
            )
        if spec["id"] == "nitro":
            protocol["composicaoDelta"] = extract_delta_composition(ws)

        protocols.append(protocol)

    total_treatments = sum(protocol["totalTratamentos"] for protocol in protocols)
    data = {
        "meta": {
            "titulo": "Caderno de Campo - Protocolos de Pesquisa",
            "parceiros": "Excelencia / Terram",
            "safra": "2026",
            "totalProtocolos": len(protocols),
            "totalTratamentos": total_treatments,
            "repeticoes": 4,
            "fontes": [WORKBOOK.name, FIELD_BOOK.name],
            "avisos": [
                "As receitas foram mantidas como escritas nos arquivos de origem, inclusive abreviacoes e grafias a revisar.",
                "A area marcada como XXXXX no croqui original foi identificada no app como area nao identificada.",
            ],
        },
        "croquiGeral": GENERAL_MAP,
        "protocolos": protocols,
    }

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    payload = json.dumps(data, ensure_ascii=False, separators=(",", ":"))
    OUTPUT.write_text(f"window.EXA_PROTOCOL_DATA = {payload};\n", encoding="utf-8")
    print(f"Generated {OUTPUT} with {len(protocols)} protocols and {total_treatments} treatments")


if __name__ == "__main__":
    main()
